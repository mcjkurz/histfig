"""
Document processing utilities for extracting text from PDFs, text files, and DOCX files,
and chunking them for vector storage.
"""

import os
import re
from typing import List, Dict, Any
import PyPDF2
import logging
from io import BytesIO
import nltk
import jieba
from docx import Document
from config import CHUNK_SIZE, MAX_CHUNK_CHARS, OVERLAP_PERCENT

class DocumentProcessor:
    def __init__(self, chunk_size: int = None, max_chunk_chars: int = None,
                 overlap_percent: int = None):
        """
        Initialize document processor with character-based and word-based chunking.
        
        Args:
            chunk_size: Words per chunk (defaults to config CHUNK_SIZE)
            max_chunk_chars: Characters per chunk (defaults to config MAX_CHUNK_CHARS)
            overlap_percent: Overlap percentage 0-50% (defaults to config OVERLAP_PERCENT)
        """
        # Use config defaults if not provided
        self.chunk_size = chunk_size if chunk_size is not None else CHUNK_SIZE
        self.max_chunk_chars = max_chunk_chars if max_chunk_chars is not None else MAX_CHUNK_CHARS
        self.overlap_percent = max(0, min(50, overlap_percent if overlap_percent is not None else OVERLAP_PERCENT))
        
        # Calculate overlaps from percentage
        self.chunk_overlap = int(self.chunk_size * self.overlap_percent / 100)
        self.char_overlap = int(self.max_chunk_chars * self.overlap_percent / 100)
        
        # Download NLTK data for sentence tokenization
        try:
            nltk.data.find('tokenizers/punkt')
        except LookupError:
            try:
                nltk.download('punkt', quiet=True)
            except Exception as e:
                logging.warning(f"Failed to download NLTK punkt data: {e}")
    
    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """
        Extract text from PDF file content.
        
        Args:
            file_content: PDF file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            pdf_file = BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text += page_text
                except Exception as e:
                    logging.warning(f"Error extracting text from page: {e}")
                    continue
            
            return text
        except Exception as e:
            logging.error(f"Error extracting text from PDF: {e}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    def extract_text_from_txt(self, file_content: bytes) -> str:
        """
        Extract text from text file content.
        
        Args:
            file_content: Text file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    return text
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            text = file_content.decode('utf-8', errors='replace')
            return text
            
        except Exception as e:
            logging.error(f"Error extracting text from text file: {e}")
            raise Exception(f"Failed to extract text from text file: {str(e)}")
    
    def extract_text_from_docx(self, file_content: bytes) -> str:
        """
        Extract text from DOCX file content.
        
        Args:
            file_content: DOCX file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            docx_file = BytesIO(file_content)
            doc = Document(docx_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                paragraph_text = paragraph.text.strip()
                if paragraph_text:
                    text += paragraph_text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            text += cell_text + " "
                text += "\n"
            
            return text
        except Exception as e:
            logging.error(f"Error extracting text from DOCX: {e}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    
    def clean_text(self, text: str) -> str:
        """
        Clean and normalize text by keeping only alphabetic characters, spaces, and punctuation.
        
        Args:
            text: Raw text
            
        Returns:
            Cleaned text with only alphabet, spaces, and punctuation
        """
        # Remove all newlines and normalize whitespace
        text = text.replace('\n', ' ').replace('\r', ' ')
        text = re.sub(r'\s+', ' ', text)
        
        # Keep only alphabetic characters, numbers, spaces, and punctuation
        text = re.sub(r'[^\w\s\.,;:!?\'"()\-–—\[\]{}]', '', text, flags=re.UNICODE)
        
        return text.strip()
    
    def chunk_text(self, text: str, metadata: Dict[str, Any] = None) -> List[Dict[str, Any]]:
        """
        Split text into chunks with intelligent boundary detection.
        Uses character-based chunking with configurable size to avoid splitting mid-word/sentence.
        
        Args:
            text: Text to chunk
            metadata: Base metadata for all chunks
            
        Returns:
            List of text chunks with metadata
        """
        if metadata is None:
            metadata = {}
        
        # Clean text minimally - just normalize whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        
        # Use configured character-based chunk size
        target_chunk_size = self.max_chunk_chars
        overlap_size = self.char_overlap
        
        # If text is small, return as single chunk
        if len(text) <= target_chunk_size:
            return [{
                "text": text,
                "metadata": {
                    **metadata, 
                    "chunk_index": 0, 
                    "total_chunks": 1,
                    "start_char": 0,
                    "end_char": len(text),
                    "char_count": len(text)
                }
            }]
        
        chunks = []
        chunk_index = 0
        start_pos = 0
        
        while start_pos < len(text):
            # Calculate end position
            end_pos = start_pos + target_chunk_size
            
            # If this is the last chunk, just take the rest
            if end_pos >= len(text):
                chunk_text = text[start_pos:]
                chunks.append({
                    "text": chunk_text,
                    "metadata": {
                        **metadata,
                        "chunk_index": chunk_index,
                        "start_char": start_pos,
                        "end_char": len(text),
                        "char_count": len(chunk_text)
                    }
                })
                break
            
            # Find a good break point near end_pos to avoid splitting mid-word
            # Look for punctuation or whitespace within 50 chars before end_pos
            search_start = max(start_pos, end_pos - 50)
            
            # Chinese and English punctuation marks that indicate good break points
            break_chars = '。！？；.!?;\n '
            
            # Search backwards from end_pos for a good break point
            best_break = end_pos
            for i in range(end_pos - 1, search_start - 1, -1):
                if text[i] in break_chars:
                    best_break = i + 1  # Break after the punctuation/space
                    break
            
            # Extract chunk
            chunk_text = text[start_pos:best_break].strip()
            
            if chunk_text:  # Only add non-empty chunks
                chunks.append({
                    "text": chunk_text,
                    "metadata": {
                        **metadata,
                        "chunk_index": chunk_index,
                        "start_char": start_pos,
                        "end_char": best_break,
                        "char_count": len(chunk_text)
                    }
                })
                chunk_index += 1
            
            # Move to next chunk with overlap
            # Calculate target overlap position
            target_overlap_start = best_break - overlap_size
            
            # Find a good break point for the overlap start (avoid mid-word)
            # Search forward from target position to find a space or punctuation
            overlap_break_chars = ' \n。！？；.!?;'
            overlap_start = target_overlap_start
            
            # Search within 50 chars forward to find a good start point
            search_end = min(target_overlap_start + 50, best_break)
            for i in range(target_overlap_start, search_end):
                if text[i] in overlap_break_chars:
                    overlap_start = i + 1  # Start after the break char
                    break
            
            start_pos = overlap_start
            
            # Make sure we don't go backwards or stay at the same position
            if start_pos <= chunks[-1]["metadata"]["start_char"]:
                start_pos = best_break
        
        # Update total_chunks in all metadata
        total_chunks = len(chunks)
        for chunk in chunks:
            chunk["metadata"]["total_chunks"] = total_chunks
        
        return chunks
    
    
    def process_file(self, file_content: bytes, filename: str, file_type: str) -> List[Dict[str, Any]]:
        """
        Process a file and return text chunks with metadata.
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            file_type: File type ('pdf', 'txt', or 'docx')
            
        Returns:
            List of processed text chunks
        """
        try:
            # Extract text based on file type
            if file_type.lower() == 'pdf':
                text = self.extract_text_from_pdf(file_content)
            elif file_type.lower() in ['txt', 'text']:
                text = self.extract_text_from_txt(file_content)
            elif file_type.lower() == 'docx':
                text = self.extract_text_from_docx(file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            if not text.strip():
                raise ValueError("No text content found in file")
            
            # Create base metadata
            base_metadata = {
                "filename": filename,
                "file_type": file_type,
                "file_size": len(file_content),
                "text_length": len(text)
            }
            
            # Chunk the text
            chunks = self.chunk_text(text, base_metadata)
            
            logging.info(f"Processed {filename}: {len(chunks)} chunks created")
            return chunks
            
        except Exception as e:
            logging.error(f"Error processing file {filename}: {e}")
            raise Exception(f"Failed to process {filename}: {str(e)}")

# Global instance
document_processor = DocumentProcessor()
